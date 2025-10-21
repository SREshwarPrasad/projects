# backend/report_generator.py
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

def generate_report_pdf_bytes(data: dict) -> bytes:
    """
    Simple PDF generator using reportlab. Returns raw bytes.
    """
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, height - 60, "Crop Yield Prediction Report")
    c.setFont("Helvetica", 12)
    y = height - 100
    for k, v in data.items():
        c.drawString(40, y, f"{k}: {v}")
        y -= 20
        if y < 80:
            c.showPage()
            y = height - 60
    c.showPage()
    c.save()
    buf.seek(0)
    return buf.read()

def generate_report_docx_bytes(data: dict) -> bytes:
    doc = Document()
    doc.add_heading("Crop Yield Prediction Report", level=1)
    for k, v in data.items():
        doc.add_paragraph(f"{k}: {v}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
